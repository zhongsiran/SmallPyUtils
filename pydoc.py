#! python3
# code:'utf8'

import docx
from docx.shared import Pt
import tkinter as tk  # 用于打开文件窗口
from tkinter import filedialog  # 用于打开文件窗口
import pyperclip
import re

# 以下用于打开文件窗口
root = tk.Tk()
root.withdraw()
file_path = filedialog.askopenfilename()

# 初始化变量
fulltext = ''
hearing_text = ''
judge_text = ''
block_start_index = 1

content = {}
if 'docx' in file_path or 'DOCX'in file_path:  # 判断文件名
    doc = docx.Document(file_path)

    # 取得终结报告的内容
    search_start_index = 3  # 略去开头两面，即标题的内容
    for current_paragraph_index in range(0, len(doc.paragraphs)):  # 从文档开头开始找第一个“请领导审批”
        if '请领导审批' in doc.paragraphs[current_paragraph_index].text:  # 当在第i段找到后：
            full_text_list = []
            for j in range(search_start_index, current_paragraph_index + 1):
                full_text_list.append(doc.paragraphs[j].text)
                fulltext = "\n  ".join(full_text_list)
                fulltext = "  " + fulltext
                search_start_index = current_paragraph_index + 1  # 下次从第i+1段开始找
            content['zhong_jie_bao_gao'] = fulltext  # 将终结报告内容存入
            
    # 取得听证告知书的内容
    for current_paragraph_index in range(search_start_index, len(doc.paragraphs)):
        if '由本局立案调查' in doc.paragraphs[current_paragraph_index].text:
            block_start_index = current_paragraph_index
            ting_zheng_number_index = current_paragraph_index - 1
        elif '视为放弃此权利' in doc.paragraphs[current_paragraph_index].text:
                hearing_list = []
                for m in range(block_start_index, current_paragraph_index + 1):
                    hearing_list.append(doc.paragraphs[m].text)
                    hearing_text = "\n  ".join(hearing_list)
                    hearing_text = "  " + hearing_text
                    search_start_index = current_paragraph_index + 1
                content['ting_zheng_gao_zhi'] = hearing_text
                
    # 取得处罚决定书的内容
    for current_paragraph_index in range(search_start_index, len(doc.paragraphs)):
        if '行政处罚决定书' in doc.paragraphs[current_paragraph_index].text:
            block_start_index = current_paragraph_index + 2
            jue_ding_shu_hao_index = current_paragraph_index + 1
        elif '六个月内直接向人民法院起诉' in doc.paragraphs[current_paragraph_index].text:
                judge_list=[]
                for m in range(block_start_index, current_paragraph_index + 1):
                    judge_list.append(doc.paragraphs[m].text)
                    judge_text = "\n  ".join(judge_list)
                    judge_text = "  " + judge_text
                    search_start_index= current_paragraph_index + 1
                content['chu_fa_jue_ding'] = judge_text

    # 取得行政处罚建议的表格内容
    xing_chu_jian_yi = doc.tables[0].rows[2].cells[1].text  # 第1个表格里面
    xing_chu_jian_yi = re.sub(r'(?<=呈领导审批)\D*\d+年\d+月\d+日', "", xing_chu_jian_yi)
    content['xing_chu_jian_yi'] = xing_chu_jian_yi

    # 第二个表格是告知书的送达回证

    # 取得处罚决定审批表
    xing_chu_jue_ding = doc.tables[2].rows[2].cells[1].text  # 第3个表格里面
    xing_chu_jue_ding = re.sub(r'(?<=呈领导审批)\D*\d+年\d+月\d+日', "", xing_chu_jue_ding)
    content['xing_chu_jue_ding'] = xing_chu_jue_ding
else:
    print('请选择DOCX文件')
    exit()
# 以下选择剪贴板内容
choice = ''
while 'q' not in choice:
    print('请输入关键字(zj,tz,jy,jd,sp)')
    choice = input()
    if choice == 'zj':
        pyperclip.copy(content['zhong_jie_bao_gao'])
        print('终结报告')
    if choice == 'tz':
        pyperclip.copy(content['ting_zheng_gao_zhi'])
        print('听证告知书')
    if choice == 'jy':
        pyperclip.copy(content['xing_chu_jian_yi'])
        print('行政处罚建议审批表')
    if choice == 'jd':
        pyperclip.copy(content['chu_fa_jue_ding'])
        print('行政处罚决定书')
    if choice =='sp':
        pyperclip.copy(content['xing_chu_jue_ding'])
        print('行政处罚决定审批表')
    if choice == 'jdh':
        #  replace word
        jue_ding_number = input('输入新的决定书号，输入c或C即从剪贴板读取')
        if jue_ding_number == 'c' or jue_ding_number == 'C':
            jue_ding_number = pyperclip.paste()
        print(jue_ding_number)
        old_jue_ding_hao = doc.paragraphs[jue_ding_shu_hao_index].text
        new_jue_ding_hao = re.sub(r'(?<=〕)\d*号', jue_ding_number + '号', old_jue_ding_hao)
        print('new:' + new_jue_ding_hao)
        doc.paragraphs[jue_ding_shu_hao_index].text = new_jue_ding_hao
        doc.paragraphs[jue_ding_shu_hao_index].style.font.name = '仿宋'
        doc.paragraphs[jue_ding_shu_hao_index].style.font.size = Pt(16)
        try:
            doc.save(file_path)
        except PermissionError:
            print('已经打开该文件，无法保存')